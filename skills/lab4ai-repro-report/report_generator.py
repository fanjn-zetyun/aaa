import os
import sys
import subprocess

# 极客级自愈：如果服务器没装 python-docx，自动静默安装
try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

def generate_report(repo_name, project_profile, implementation_steps, results_comparison, optimization_suggestions, font_english='Times New Roman', font_chinese='微软雅黑'):
    try:
        doc = docx.Document()
        
        # --- 全局字体设置 ---
        for style in doc.styles:
            if hasattr(style, 'font'):
                style.font.name = font_english
                # 修改 Word 底层的东亚字体属性
                if style._element.rPr is not None:
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_chinese)
        
        # --- 标题 ---
        title = doc.add_heading(f'【复现报告】{repo_name} - 自动化复现报告', 0)
        title.alignment = 1 # 居中
        
        # --- 1. 项目档案 ---
        doc.add_heading('一、 项目档案', level=1)
        doc.add_paragraph(project_profile)
        
        # --- 2. 复现实施步骤 ---
        doc.add_heading('二、 复现实施步骤', level=1)
        steps_map =[
            ('2.1 代码获取', 'code_fetch'),
            ('2.2 环境搭建与排坑记录', 'env_setup'),
            ('2.3 数据与参数配置', 'data_params'),
            ('2.4 训练/推理核心流程', 'core_loop'),
            ('2.5 评估流程', 'eval_process')
        ]
        for title_str, key in steps_map:
            doc.add_heading(title_str, level=2)
            content = implementation_steps.get(key, '未提供信息')
            doc.add_paragraph(str(content))
            
        # --- 3. 结果对比 ---
        doc.add_heading('三、 结果对比 (原论文 vs 当前复现)', level=1)
        if results_comparison and len(results_comparison) > 0:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '评估维度/指标'
            hdr_cells[1].text = '官方/原论文基准'
            hdr_cells[2].text = '本次实际复现值'
            
            # 设置表头加粗
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for item in results_comparison:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('metric_name', '-'))
                row_cells[1].text = str(item.get('official_value', '-'))
                row_cells[2].text = str(item.get('reproduced_value', '-'))
        else:
            doc.add_paragraph("⚠️ 本次复现未捕获到可用于对比的量化指标数据。")
            
        # --- 4. 优化建议 ---
        doc.add_heading('四、 后期全量训练与优化建议', level=1)
        doc.add_paragraph(optimization_suggestions)
        
        # --- 自动落盘保存 ---
        save_dir = f"/root/.openclaw/workspace/{repo_name}"
        os.makedirs(save_dir, exist_ok=True)
        save_path = os.path.join(save_dir, f"{repo_name}_Final_Repro_Report.docx")
        
        doc.save(save_path)
        return f"✅ 报告排版成功！已生成极度详尽的 Word 报告并保存至绝对路径：\n`{save_path}`"
        
    except Exception as e:
        return f"❌ 报告生成遭遇底层失败：{str(e)}"