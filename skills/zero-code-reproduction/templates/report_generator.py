"""
Zero-Code Reproduction Report Generator
========================================
接收 JSON 配置，生成标准格式 .docx 复现报告。

用法:
  python report_generator.py --config report_config.json --output report.docx

JSON 配置 schema:
{
  "project_name": "GeneCLR",
  "paper_title": "Protein and Genomic Language Models ...",
  "paper_venue": "Science 2026",
  "paper_doi": "10.1126/science.adv8275",
  "paper_github": "https://github.com/...",        // 可选, 无则写 "无公开代码"
  "repro_level": "Level 2 轻量验证",                // Level 2 或 Level 3
  "repro_platform": "Lab4AI (CPU + H100 80GB GPU)",
  "routing": {
    "domain": "BIOINFO",
    "type": "HYBRID",
    "confidence": 0.95,
    "plugins": ["zero-code-repro-biodefense", "zero-code-repro-csai"]
  },
  "paper_summary": "项目简介文字...",
  "architecture": "核心架构描述...",
  "paper_metrics": [
    {"name": "Micro-AUROC", "value": "0.97", "note": "65类分类"},
    ...
  ],
  "strategy": "复现策略描述...",
  "steps": [
    {"id": "Step 0", "name": "远程实例初始化", "location": "远程 CPU", "status": "✅"},
    ...
  ],
  "scaffold_files": [
    {"category": "模型", "file": "albert_df.py", "desc": "ALBERT-DF 距离感知模型"},
    ...
  ],
  "training_results": [
    {"test": "ALBERT Forward+Backward", "result": "✅ PASS", "loss_start": "7.1363", "loss_end": "7.1186", "steps": "20", "note": "MLM loss"},
    ...
  ],
  "model_params": [
    {"model": "ALBERT-DF", "repro_params": "3.3M", "paper_params": "44M", "note": "小配置验证"},
    ...
  ],
  "env_config": [
    {"key": "GPU", "value": "NVIDIA H100 80GB HBM3"},
    {"key": "PyTorch", "value": "2.6.0+cu124"},
    ...
  ],
  "pitfalls": [
    {"problem": "SSH heredoc引号冲突", "cause": "三层引号嵌套", "solution": "SCP脚本上传→SSH执行"},
    ...
  ],
  "compute_cost": [
    {"phase": "Step 0-8", "instance": "CPU 2C/4G", "duration": "3h44m", "unit_price": "¥0.3/h", "cost": "≈ ¥1.12"},
    ...
  ],
  "compute_total": "≈ ¥4.87",
  "compute_note": "本次为 Level 2 轻量验证...",      // 可选
  "confidence": [
    {"module": "ALBERT-DF 架构", "level": "HIGH", "note": "论文描述详细"},
    ...
  ],
  "todos": [                                         // 可选
    "6D距离向量第5/6维需与原作者确认",
    ...
  ],
  "conclusion": "结论正文...",
  "conclusion_results": "核心验证结果...",
  "conclusion_limitations": "局限性..."              // 可选
}
"""

import argparse
import json
import os
import sys
import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ============================================================
# Helpers
# ============================================================

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    return h


def add_table(doc, headers, rows):
    """Add a styled table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_kv_table(doc, pairs):
    """Add a 2-column key-value table."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(pairs):
        table.rows[i].cells[0].text = str(k)
        table.rows[i].cells[1].text = str(v)
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# Report Builder
# ============================================================

def build_report(cfg, output_path):
    doc = Document()

    # -- Global style --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    project = cfg.get("project_name", "Unknown")
    today = datetime.date.today().isoformat()

    # ========== Title Page ==========
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{project} 零代码复现报告')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(cfg.get("paper_title", ""))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    venue = cfg.get("paper_venue", "")
    doi = cfg.get("paper_doi", "")
    level = cfg.get("repro_level", "Level 2")
    platform = cfg.get("repro_platform", "Lab4AI")
    meta.add_run(f'{venue} · DOI: {doi}\n').font.size = Pt(11)
    meta.add_run(f'报告生成日期: {today}\n').font.size = Pt(11)
    meta.add_run(f'复现方式: Zero-Code Reproduction Scaffolding ({level})\n').font.size = Pt(11)
    meta.add_run(f'复现平台: {platform}').font.size = Pt(11)

    doc.add_page_break()

    # ========== 1. 项目概述 ==========
    add_heading_styled(doc, '1. 项目概述')

    if cfg.get("paper_summary"):
        doc.add_paragraph(cfg["paper_summary"])

    if cfg.get("architecture"):
        add_heading_styled(doc, '1.1 核心架构', level=2)
        doc.add_paragraph(cfg["architecture"])

    if cfg.get("paper_metrics"):
        add_heading_styled(doc, '1.2 论文关键指标', level=2)
        add_table(doc,
                  ['指标', '论文报告值', '说明'],
                  [[m["name"], m["value"], m.get("note", "")] for m in cfg["paper_metrics"]])

    doc.add_page_break()

    # ========== 2. 复现方法 ==========
    add_heading_styled(doc, '2. 复现方法')

    if cfg.get("strategy"):
        add_heading_styled(doc, '2.1 复现策略', level=2)
        doc.add_paragraph(cfg["strategy"])

    if cfg.get("steps"):
        add_heading_styled(doc, '2.2 零代码复现流程', level=2)
        add_table(doc,
                  ['步骤', '内容', '执行位置', '状态'],
                  [[s["id"], s["name"], s.get("location", ""), s.get("status", "⏳")]
                   for s in cfg["steps"]])

    doc.add_page_break()

    # ========== 3. 代码脚手架产出 ==========
    if cfg.get("scaffold_files"):
        add_heading_styled(doc, '3. 代码脚手架产出')

        add_heading_styled(doc, f'3.1 生成文件清单 ({len(cfg["scaffold_files"])} 文件)', level=2)
        add_table(doc,
                  ['类别', '文件', '说明'],
                  [[f["category"], f["file"], f.get("desc", "")] for f in cfg["scaffold_files"]])

        if cfg.get("routing"):
            r = cfg["routing"]
            add_heading_styled(doc, '3.2 学科路由结果', level=2)
            add_kv_table(doc, [
                ('Domain', r.get("domain", "")),
                ('Type', r.get("type", "")),
                ('Confidence', str(r.get("confidence", ""))),
                ('激活插件', ', '.join(r.get("plugins", []))),
            ])

        doc.add_page_break()

    # ========== 4. 训练验证结果 ==========
    if cfg.get("training_results"):
        add_heading_styled(doc, '4. GPU 训练验证结果')

        add_heading_styled(doc, '4.1 验证项目', level=2)
        add_table(doc,
                  ['测试项', '结果', 'Loss 起始', 'Loss 终止', '步数', '说明'],
                  [[t["test"], t["result"], t.get("loss_start", "-"),
                    t.get("loss_end", "-"), t.get("steps", "-"), t.get("note", "")]
                   for t in cfg["training_results"]])

    if cfg.get("model_params"):
        add_heading_styled(doc, '4.2 模型参数统计', level=2)
        add_table(doc,
                  ['模型', '本次复现参数量', '论文参数量', '说明'],
                  [[m["model"], m["repro_params"], m["paper_params"], m.get("note", "")]
                   for m in cfg["model_params"]])

    if cfg.get("env_config"):
        add_heading_styled(doc, '4.3 环境配置', level=2)
        add_kv_table(doc, [(e["key"], e["value"]) for e in cfg["env_config"]])

    doc.add_page_break()

    # ========== 5. 排坑记录 ==========
    if cfg.get("pitfalls"):
        add_heading_styled(doc, '5. 排坑记录')
        add_table(doc,
                  ['问题', '原因', '解决方案'],
                  [[p["problem"], p["cause"], p["solution"]] for p in cfg["pitfalls"]])
        doc.add_page_break()

    # ========== 6. 算力消耗 ==========
    if cfg.get("compute_cost"):
        add_heading_styled(doc, '6. 算力消耗')
        rows = [[c["phase"], c["instance"], c["duration"], c["unit_price"], c["cost"]]
                for c in cfg["compute_cost"]]
        if cfg.get("compute_total"):
            rows.append(['总计', '-', '-', '-', cfg["compute_total"]])
        add_table(doc, ['阶段', '实例类型', '运行时长', '单价', '费用'], rows)

        if cfg.get("compute_note"):
            doc.add_paragraph()
            doc.add_paragraph(cfg["compute_note"])

    # ========== 7. 置信度评估 ==========
    if cfg.get("confidence"):
        add_heading_styled(doc, '7. 置信度评估')
        add_table(doc,
                  ['模块', '置信度', '说明'],
                  [[c["module"], c["level"], c.get("note", "")] for c in cfg["confidence"]])

        if cfg.get("todos"):
            doc.add_paragraph()
            doc.add_paragraph('TODO (需人工补全):\n' + '\n'.join(f'• {t}' for t in cfg["todos"]))

    # ========== 8. 结论 ==========
    add_heading_styled(doc, '8. 结论')

    if cfg.get("conclusion"):
        doc.add_paragraph(cfg["conclusion"])
    if cfg.get("conclusion_results"):
        doc.add_paragraph(cfg["conclusion_results"])
    if cfg.get("conclusion_limitations"):
        doc.add_paragraph(cfg["conclusion_limitations"])

    # ========== Save ==========
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"Report saved: {output_path} ({size_kb:.1f} KB)")
    return output_path


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='Zero-Code Reproduction Report Generator')
    parser.add_argument('--config', required=True, help='Path to JSON config file')
    parser.add_argument('--output', required=True, help='Output .docx path')
    args = parser.parse_args()

    with open(args.config, 'r', encoding='utf-8') as f:
        cfg = json.load(f)

    build_report(cfg, args.output)


if __name__ == '__main__':
    main()
