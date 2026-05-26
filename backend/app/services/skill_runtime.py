"""Runtime adapters for executable skills.

The skill files stay as local task definitions. This module reads their tool
metadata and provides backend-owned adapters for the executable parts so Web
workflow code never has to run shell snippets from prompts directly.
"""

from __future__ import annotations

import asyncio
import importlib.util
import json
import os
import re
import shutil
import time
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml


@dataclass(frozen=True, slots=True)
class SkillToolSpec:
    name: str
    description: str
    entry_point: str
    input_schema: dict[str, Any]
    skill_name: str
    base_dir: Path
    aliases: tuple[str, ...] = ()


@dataclass(slots=True)
class SkillRuntimeResult:
    name: str
    content: str
    ok: bool = True
    metadata: dict[str, Any] = field(default_factory=dict)


class SkillRuntime:
    def __init__(self, skills_dir: Path, workspace_root: Path) -> None:
        self.skills_dir = skills_dir
        self.workspace_root = workspace_root
        self._specs = self._load_specs()

    def list_specs(self) -> list[SkillToolSpec]:
        return sorted(self._specs.values(), key=lambda item: item.name)

    def spec(self, name: str) -> SkillToolSpec | None:
        return self._specs.get(name) or next(
            (spec for spec in self._specs.values() if name in spec.aliases),
            None,
        )

    def health(self) -> list[dict[str, Any]]:
        checks: list[dict[str, Any]] = []
        for spec in self.list_specs():
            adapter = _adapter_for(spec.name)
            status = {
                "name": spec.name,
                "skill_name": spec.skill_name,
                "entry_point": spec.entry_point,
                "aliases": list(spec.aliases),
                "adapter": adapter is not None,
                "loaded": False,
                "error": None,
            }
            try:
                if adapter and spec.name == "paper_analyze":
                    # The bundled skill metadata points at a missing wrapper.
                    # The backend adapter uses the script's real parsing
                    # functions without modifying the skill directory.
                    status["loaded"] = True
                    status["entrypoint_note"] = "backend_adapter_wraps_existing_script_functions"
                else:
                    _load_entrypoint(spec.base_dir, spec.entry_point)
                    status["loaded"] = True
            except Exception as exc:  # pragma: no cover - exact import errors are env-specific
                status["error"] = f"{type(exc).__name__}: {exc}"
            checks.append(status)
        return checks

    async def invoke(self, name: str, payload: dict[str, Any]) -> SkillRuntimeResult:
        spec = self.spec(name)
        if spec is None:
            return SkillRuntimeResult(
                name,
                f"未找到可执行 skill tool：{name}",
                ok=False,
                metadata={"error_code": "unknown_skill_tool"},
            )

        adapter = _adapter_for(spec.name)
        if adapter is None:
            return SkillRuntimeResult(
                spec.name,
                f"skill tool {spec.name} 暂无后端执行适配器。",
                ok=False,
                metadata={"error_code": "missing_adapter", "entry_point": spec.entry_point},
            )
        return await adapter(self, spec, payload)

    def _load_specs(self) -> dict[str, SkillToolSpec]:
        specs: dict[str, SkillToolSpec] = {}
        if not self.skills_dir.exists():
            return specs

        for skill_dir in sorted(path for path in self.skills_dir.iterdir() if path.is_dir()):
            tools_file = skill_dir / "tools.yaml"
            if tools_file.exists():
                for raw in _read_yaml(tools_file, default=[]):
                    if isinstance(raw, dict) and raw.get("name"):
                        spec = _spec_from_mapping(raw, skill_dir)
                        specs[spec.name] = spec

            manifest_file = skill_dir / "manifest.yaml"
            if manifest_file.exists():
                raw_manifest = _read_yaml(manifest_file, default={})
                if isinstance(raw_manifest, dict) and raw_manifest.get("name"):
                    spec = _spec_from_mapping(raw_manifest, skill_dir)
                    specs[spec.name] = spec

        return specs


async def _run_repo_audit(
    runtime: SkillRuntime,
    spec: SkillToolSpec,
    payload: dict[str, Any],
) -> SkillRuntimeResult:
    repo_url = str(payload.get("repo_url") or payload.get("github_url") or "").strip()
    if not repo_url:
        return SkillRuntimeResult(
            "repo_audit",
            "未提供 GitHub 仓库 URL，无法执行仓库审计。",
            ok=False,
            metadata={"error_code": "missing_repo_url"},
        )

    clone_url = _github_clone_url(repo_url, payload)
    output_dir = _safe_output_dir(runtime, payload, repo_url=repo_url)
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        module = _load_module(spec.base_dir / "scripts" / "main.py")
        audit_repo = getattr(module, "audit_repo")
        proxy_url = _proxy_url_from_payload(payload)
        _set_module_proxy(module, proxy_url)
        report = await asyncio.to_thread(
            _invoke_with_repo_audit_env,
            module,
            audit_repo,
            clone_url,
            proxy_url=proxy_url,
        )
    except Exception as exc:
        return SkillRuntimeResult(
            "repo_audit",
            f"仓库审计执行失败：{type(exc).__name__}: {exc}",
            ok=False,
            metadata={"error_code": "repo_audit_failed"},
        )

    report_path = output_dir / "repo_audit.md"
    report_path.write_text(str(report), encoding="utf-8")
    score = _parse_score(str(report))
    ok = not str(report).lstrip().startswith("❌")
    return SkillRuntimeResult(
        "repo_audit",
        str(report),
        ok=ok,
        metadata={
            "repo_url": repo_url,
            "clone_url": clone_url,
            "repo": _repo_slug(repo_url),
            "repo_name": _repo_name(repo_url),
            "score": score,
            "report_path": str(report_path),
            "artifact_paths": [str(report_path)],
            "proxy_used": bool(proxy_url),
            "error_code": None if ok else "repo_audit_failed",
        },
    )


async def _run_paper_analyze(
    runtime: SkillRuntime,
    spec: SkillToolSpec,
    payload: dict[str, Any],
) -> SkillRuntimeResult:
    paper_url = str(payload.get("paper_url") or "").strip()
    paper_path = str(payload.get("paper_path") or "").strip()
    github_url = str(payload.get("github_url") or "").strip()
    if not paper_url and not paper_path:
        return SkillRuntimeResult(
            "paper_analyze",
            "未提供论文 PDF URL 或本地 PDF 路径，跳过论文分析。",
            ok=False,
            metadata={"error_code": "missing_paper_source"},
        )

    output_dir = _safe_output_dir(runtime, payload, repo_url=github_url or paper_url)
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        module = _load_module(spec.base_dir / "scripts" / "analyze_paper.py")
        pdf_path = Path(paper_path) if paper_path else output_dir / "paper.pdf"
        full_text = ""
        pages = 0
        if paper_url and not paper_path:
            ok = await asyncio.to_thread(_download_pdf, paper_url, pdf_path)
            if not ok:
                full_text = await asyncio.to_thread(_download_arxiv_abs_text, paper_url)
                if not full_text:
                    return SkillRuntimeResult(
                        "paper_analyze",
                        f"论文 PDF 下载失败：{paper_url}",
                        ok=False,
                        metadata={"error_code": "paper_download_failed", "paper_url": paper_url},
                    )
                pdf_path = output_dir / "paper_abs.html"
                pdf_path.write_text(full_text, encoding="utf-8")
        if not full_text:
            full_text, pages = await asyncio.to_thread(module.extract_text_from_pdf, str(pdf_path))
        text_path = output_dir / "paper_text.txt"
        text_path.write_text(full_text, encoding="utf-8")
        info = await asyncio.to_thread(module.parse_paper_text, full_text)
        paper_name = str(payload.get("paper_name") or module.infer_paper_name(github_url or paper_url))
        info["paper_name"] = paper_name
        report_md = module.build_markdown_report(paper_name, info)
        report_path = output_dir / "paper_analysis.md"
        report_path.write_text(report_md, encoding="utf-8")
    except Exception as exc:
        return SkillRuntimeResult(
            "paper_analyze",
            f"论文分析执行失败：{type(exc).__name__}: {exc}",
            ok=False,
            metadata={"error_code": "paper_analyze_failed"},
        )

    return SkillRuntimeResult(
        "paper_analyze",
        report_md,
        metadata={
            "paper_url": paper_url,
            "paper_path": str(pdf_path),
            "paper_name": paper_name,
            "pages": pages,
            "score": info.get("score", 0),
            "arxiv_id": info.get("arxiv_id", ""),
            "venue": info.get("venue", ""),
            "metrics": info.get("metrics", {}),
            "hyperparams": info.get("hyperparams", {}),
            "datasets": info.get("datasets", []),
            "innovation_points": info.get("innovation_points", []),
            "tables": info.get("tables", {}),
            "report_path": str(report_path),
            "text_path": str(text_path),
            "artifact_paths": [str(report_path), str(text_path)],
        },
    )


async def _run_repro_report(
    runtime: SkillRuntime,
    spec: SkillToolSpec,
    payload: dict[str, Any],
) -> SkillRuntimeResult:
    repo_name = re.sub(
        r"[^A-Za-z0-9_.-]+",
        "-",
        str(payload.get("repo_name") or "project"),
    ).strip("-") or "project"
    project_profile = str(payload.get("project_profile") or "未提供项目档案。")
    implementation_steps = payload.get("implementation_steps")
    if not isinstance(implementation_steps, dict):
        implementation_steps = {}
    results_comparison = payload.get("results_comparison")
    if not isinstance(results_comparison, list):
        results_comparison = []
    optimization_suggestions = str(payload.get("optimization_suggestions") or "暂无。")

    try:
        generate_report = _load_entrypoint(spec.base_dir, spec.entry_point)
        result_text = await asyncio.to_thread(
            generate_report,
            repo_name=repo_name,
            project_profile=project_profile,
            implementation_steps=implementation_steps,
            results_comparison=results_comparison,
            optimization_suggestions=optimization_suggestions,
            font_english=str(payload.get("font_english") or "Times New Roman"),
            font_chinese=str(payload.get("font_chinese") or "微软雅黑"),
        )
    except Exception as exc:
        return SkillRuntimeResult(
            "generate_repro_report",
            f"复现报告生成失败：{type(exc).__name__}: {exc}",
            ok=False,
            metadata={"error_code": "report_generate_failed"},
        )

    generated_report_path = _extract_generated_report_path(str(result_text), repo_name)
    report_path = _workspace_repro_report_path(runtime, payload, repo_name)
    generation_source = "lab4ai-repro-report"
    if generated_report_path.exists():
        try:
            report_path.parent.mkdir(parents=True, exist_ok=True)
            if generated_report_path.resolve() != report_path.resolve():
                shutil.copyfile(generated_report_path, report_path)
        except Exception as exc:
            return SkillRuntimeResult(
                "generate_repro_report",
                f"复现报告发布到项目工作区失败：{type(exc).__name__}: {exc}",
                ok=False,
                metadata={
                    "error_code": "report_publish_failed",
                    "report_path": str(report_path),
                    "skill_output": str(result_text),
                },
            )
    else:
        generation_source = "backend_report_fallback"
        try:
            _write_backend_repro_report(
                report_path,
                repo_name=repo_name,
                project_profile=project_profile,
                implementation_steps=implementation_steps,
                results_comparison=results_comparison,
                optimization_suggestions=optimization_suggestions,
            )
        except Exception as exc:
            return SkillRuntimeResult(
                "generate_repro_report",
                f"复现报告生成后未找到 Word 文件：{report_path}",
                ok=False,
                metadata={
                    "error_code": "report_artifact_missing",
                    "report_path": str(generated_report_path),
                    "skill_output": str(result_text),
                    "fallback_error": f"{type(exc).__name__}: {exc}",
                },
            )
    return SkillRuntimeResult(
        "generate_repro_report",
        str(result_text),
        metadata={
            "repo_name": repo_name,
            "report_path": str(report_path),
            "artifact_paths": [str(report_path)],
            "generation_source": generation_source,
            "skill_output": str(result_text),
        },
    )


def _extract_generated_report_path(skill_output: str, repo_name: str) -> Path:
    match = re.search(r"`([^`]+\.docx)`", skill_output)
    if match:
        return Path(match.group(1))
    match = re.search(r"([/\w:.-]+\.docx)", skill_output)
    if match:
        return Path(match.group(1))
    return Path("/root/.openclaw/workspace") / repo_name / f"{repo_name}_Final_Repro_Report.docx"


def _workspace_repro_report_path(
    runtime: SkillRuntime,
    payload: dict[str, Any],
    repo_name: str,
) -> Path:
    conversation_id = str(payload.get("conversation_id") or "").strip()
    base_dir = runtime.workspace_root / conversation_id / repo_name if conversation_id else runtime.workspace_root / repo_name
    return base_dir / f"{repo_name}_Final_Repro_Report.docx"


def _write_backend_repro_report(
    path: Path,
    *,
    repo_name: str,
    project_profile: str,
    implementation_steps: dict[str, Any],
    results_comparison: list[Any],
    optimization_suggestions: str,
) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(f"{repo_name} Reproduction Report", level=0)
    document.add_heading("Project Profile", level=1)
    document.add_paragraph(project_profile)
    document.add_heading("Implementation Steps", level=1)
    step_titles = [
        ("Code Fetch", "code_fetch"),
        ("Environment Setup", "env_setup"),
        ("Data And Parameters", "data_params"),
        ("Core Loop", "core_loop"),
        ("Evaluation Process", "eval_process"),
    ]
    for title, key in step_titles:
        document.add_heading(title, level=2)
        document.add_paragraph(str(implementation_steps.get(key) or "Not provided."))
    document.add_heading("Results Comparison", level=1)
    if results_comparison:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Metric"
        headers[1].text = "Official"
        headers[2].text = "Reproduced"
        for item in results_comparison:
            row = table.add_row().cells
            if isinstance(item, dict):
                row[0].text = str(item.get("metric_name") or "-")
                row[1].text = str(item.get("official_value") or "-")
                row[2].text = str(item.get("reproduced_value") or "-")
            else:
                row[0].text = str(item)
                row[1].text = "-"
                row[2].text = "-"
    else:
        document.add_paragraph("No quantitative comparison data was captured.")
    document.add_heading("Optimization Suggestions", level=1)
    document.add_paragraph(optimization_suggestions)
    document.save(path)


def _write_backend_repro_report_markdown(
    path: Path,
    *,
    repo_name: str,
    project_profile: str,
    implementation_steps: dict[str, Any],
    results_comparison: list[Any],
    optimization_suggestions: str,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        f"# {repo_name} 自动化复现报告",
        "",
        "## 一、项目档案",
        "",
        _markdown_block(project_profile),
        "",
        "## 二、复现实施步骤",
        "",
    ]
    step_titles = [
        ("2.1 代码获取", "code_fetch"),
        ("2.2 环境搭建与排坑记录", "env_setup"),
        ("2.3 数据与参数配置", "data_params"),
        ("2.4 训练/推理核心流程", "core_loop"),
        ("2.5 评估流程", "eval_process"),
    ]
    for title, key in step_titles:
        lines.extend(
            [
                f"### {title}",
                "",
                _markdown_block(implementation_steps.get(key) or "未提供信息。"),
                "",
            ]
        )

    lines.extend(["## 三、结果对比 (原论文 vs 当前复现)", ""])
    if results_comparison:
        lines.extend(
            [
                "| 评估维度/指标 | 官方/原论文基准 | 本次实际复现值 |",
                "| --- | --- | --- |",
            ]
        )
        for item in results_comparison:
            if isinstance(item, dict):
                metric = item.get("metric_name") or "-"
                official = item.get("official_value") or "-"
                reproduced = item.get("reproduced_value") or "-"
            else:
                metric = item
                official = "-"
                reproduced = "-"
            lines.append(
                "| "
                + " | ".join(
                    [
                        _markdown_table_cell(metric),
                        _markdown_table_cell(official),
                        _markdown_table_cell(reproduced),
                    ]
                )
                + " |"
            )
    else:
        lines.append("本次复现未捕获到可用于对比的量化指标数据。")

    lines.extend(
        [
            "",
            "## 四、后期全量训练与优化建议",
            "",
            _markdown_block(optimization_suggestions or "暂无。"),
            "",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def _markdown_block(value: Any) -> str:
    text = str(value).strip()
    return text or "-"


def _markdown_table_cell(value: Any) -> str:
    text = str(value).replace("\n", "<br>").replace("|", "\\|").strip()
    return text or "-"


def _adapter_for(name: str):
    return {
        "repo_audit": _run_repo_audit,
        "paper_analyze": _run_paper_analyze,
        "generate_repro_report": _run_repro_report,
    }.get(name)


def _spec_from_mapping(raw: dict[str, Any], skill_dir: Path) -> SkillToolSpec:
    name = str(raw.get("name") or "").strip()
    return SkillToolSpec(
        name=name,
        description=str(raw.get("description") or ""),
        entry_point=str(raw.get("entry_point") or ""),
        input_schema=dict(raw.get("parameters") or {"type": "object", "properties": {}}),
        skill_name=skill_dir.name,
        base_dir=skill_dir,
        aliases=_aliases_for(name),
    )


def _aliases_for(name: str) -> tuple[str, ...]:
    return {
        "repo_audit": ("analyze_repo",),
        "paper_analyze": ("analyze_paper",),
        "generate_repro_report": ("repro_report",),
    }.get(name, ())


def _read_yaml(path: Path, *, default: Any) -> Any:
    raw = path.read_text(encoding="utf-8")
    try:
        data = yaml.safe_load(raw)
    except Exception:
        normalized = _normalize_legacy_yaml(raw)
        if normalized == raw:
            return default
        try:
            data = yaml.safe_load(normalized)
        except Exception:
            return default
    return default if data is None else data


def _normalize_legacy_yaml(raw: str) -> str:
    """Accept legacy skill manifests with `key:value` at line start.

    Some bundled skills contain YAML that is understood by humans but rejected
    by PyYAML, for example `required:["repo_name"]`.  Runtime adapters should
    not mutate the read-only skill directory just to recover from that typo, so
    normalize the narrow legacy shape before parsing.
    """

    return re.sub(r"(?m)^(\s*[A-Za-z_][\w-]*):(?!\s|$)", r"\1: ", raw)


def _load_entrypoint(base_dir: Path, entry_point: str):
    if ":" not in entry_point:
        raise ValueError(f"非法 entry_point：{entry_point}")
    relative_path, attr = entry_point.split(":", 1)
    module = _load_module(base_dir / relative_path)
    try:
        return getattr(module, attr)
    except AttributeError as exc:
        raise AttributeError(f"{entry_point} 缺少函数 {attr}") from exc


def _load_module(path: Path):
    module_name = "lobster_skill_" + re.sub(r"[^A-Za-z0-9_]+", "_", str(path.resolve()))
    spec = importlib.util.spec_from_file_location(module_name, path)
    if spec is None or spec.loader is None:
        raise ImportError(f"无法加载 Python 模块：{path}")
    module = importlib.util.module_from_spec(spec)
    proxy_keys = ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY")
    old_proxy_env = {key: os.environ.get(key) for key in proxy_keys}
    try:
        spec.loader.exec_module(module)
    finally:
        for key, value in old_proxy_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
    return module


def _proxy_url_from_payload(payload: dict[str, Any]) -> str | None:
    explicit = payload.get("proxy_url") or payload.get("github_proxy")
    if explicit is not None:
        value = str(explicit).strip()
        return value or None
    for key in ("LOBSTER_GITHUB_PROXY", "GITHUB_PROXY"):
        value = os.environ.get(key)
        if value and value.strip():
            return value.strip()
    return None


def _github_clone_url(repo_url: str, payload: dict[str, Any]) -> str:
    explicit = payload.get("clone_url")
    if explicit:
        return str(explicit).strip()
    prefix = str(
        payload.get("github_url_prefix")
        or os.environ.get("LOBSTER_GITHUB_URL_PREFIX")
        or os.environ.get("GITHUB_URL_PREFIX")
        or "https://gh-proxy.org/"
    ).strip()
    if not prefix:
        return repo_url
    if not _is_github_url(repo_url):
        return repo_url
    if repo_url.startswith(prefix):
        return repo_url
    return prefix.rstrip("/") + "/" + repo_url


def _set_module_proxy(module: Any, proxy_url: str | None) -> None:
    if proxy_url:
        module.PROXY_URL = proxy_url
        module.REQ_PROXIES = {"http": proxy_url, "https": proxy_url}
    else:
        module.PROXY_URL = ""
        module.REQ_PROXIES = None


def _invoke_with_proxy_env(func: Any, *args: Any, proxy_url: str | None) -> Any:
    proxy_keys = ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY")
    old_proxy_env = {key: os.environ.get(key) for key in proxy_keys}
    old_path = os.environ.get("PATH")
    try:
        for key in proxy_keys:
            os.environ.pop(key, None)
        os.environ["PATH"] = _path_with_git_usr_bin(old_path)
        if proxy_url:
            os.environ["http_proxy"] = proxy_url
            os.environ["https_proxy"] = proxy_url
        return func(*args)
    finally:
        for key, value in old_proxy_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        if old_path is None:
            os.environ.pop("PATH", None)
        else:
            os.environ["PATH"] = old_path


def _invoke_with_repo_audit_env(
    module: Any,
    func: Any,
    *args: Any,
    proxy_url: str | None,
) -> Any:
    original_run = module.subprocess.run

    def run_without_submodules(cmd: Any, *run_args: Any, **run_kwargs: Any) -> Any:
        return original_run(
            _strip_recurse_submodules_from_git_clone(cmd),
            *run_args,
            **run_kwargs,
        )

    module.subprocess.run = run_without_submodules
    try:
        return _invoke_with_proxy_env(func, *args, proxy_url=proxy_url)
    finally:
        module.subprocess.run = original_run


def _strip_recurse_submodules_from_git_clone(command: Any) -> Any:
    if not isinstance(command, (list, tuple)):
        return command
    items = list(command)
    if len(items) < 3:
        return command
    executable = str(items[0]).lower()
    if not executable.endswith("git") and not executable.endswith("git.exe"):
        return command
    if "clone" not in [str(item).lower() for item in items[1:3]]:
        return command
    stripped = [item for item in items if item != "--recurse-submodules"]
    return tuple(stripped) if isinstance(command, tuple) else stripped


def _download_pdf(
    url: str,
    output_path: Path,
    timeout: int = 120,
    max_bytes: int = 3 * 1024 * 1024,
) -> bool:
    deadline = time.monotonic() + timeout
    try:
        request = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; LobsterPaperAnalysis/1.0)"},
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with urllib.request.urlopen(request, timeout=min(timeout, 30)) as response:
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > max_bytes:
                return False
            target = output_path.open("wb")
            with target:
                total = 0
                while True:
                    if time.monotonic() > deadline:
                        raise TimeoutError(f"PDF download exceeded {timeout}s")
                    chunk = response.read(64 * 1024)
                    if not chunk:
                        break
                    total += len(chunk)
                    if total > max_bytes:
                        raise TimeoutError(f"PDF download exceeded {max_bytes} bytes")
                    target.write(chunk)
        return output_path.stat().st_size > 0
    except Exception:
        try:
            output_path.unlink(missing_ok=True)
        except OSError:
            pass
        return False


def _download_arxiv_abs_text(url: str, timeout: int = 30) -> str:
    abs_url = _arxiv_abs_url_from_pdf_url(url)
    if not abs_url:
        return ""
    try:
        request = urllib.request.Request(
            abs_url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; LobsterPaperAnalysis/1.0)"},
        )
        with urllib.request.urlopen(request, timeout=timeout) as response:
            html = response.read(512 * 1024).decode("utf-8", errors="ignore")
    except Exception:
        return ""
    return _html_to_text(html)


def _arxiv_abs_url_from_pdf_url(url: str) -> str:
    match = re.search(r"https?://arxiv\.org/(?:pdf|abs)/(\d{4}\.\d{4,5})(?:\.pdf)?", url)
    if not match:
        return ""
    return f"https://arxiv.org/abs/{match.group(1)}"


def _html_to_text(html: str) -> str:
    text = re.sub(r"<script\b.*?</script>", " ", html, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<style\b.*?</style>", " ", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _path_with_git_usr_bin(path_value: str | None) -> str:
    current = path_value or ""
    git_usr_bin = _git_usr_bin_from_git_executable(shutil.which("git"))
    if git_usr_bin is None or not git_usr_bin.exists():
        return current
    entries = [item for item in current.split(os.pathsep) if item]
    normalized = {str(Path(item)).casefold() for item in entries}
    if str(git_usr_bin).casefold() in normalized:
        return current
    return os.pathsep.join([str(git_usr_bin), *entries])


def _git_usr_bin_from_git_executable(git_executable: str | None) -> Path | None:
    if not git_executable:
        return None
    git_path = Path(git_executable)
    if git_path.parent.name.lower() == "cmd":
        return git_path.parent.parent / "usr" / "bin"
    return None


def _safe_output_dir(runtime: SkillRuntime, payload: dict[str, Any], *, repo_url: str) -> Path:
    root = runtime.workspace_root.resolve()
    conversation_id = str(payload.get("conversation_id") or "manual").strip() or "manual"
    repo_name = _repo_name(repo_url)
    requested = payload.get("output_dir")
    if requested:
        candidate = Path(str(requested))
        if not candidate.is_absolute():
            candidate = root / conversation_id / candidate
    else:
        candidate = root / conversation_id / repo_name
    resolved = candidate.resolve()
    if not _is_relative_to(resolved, root):
        raise ValueError("输出目录必须位于 runtime/workspaces 下")
    return resolved


def _repo_slug(url: str) -> str:
    url = _strip_github_proxy_prefix(url)
    if "github.com/" in url:
        return url.rstrip("/").split("github.com/", 1)[1].removesuffix(".git")
    return url.rstrip("/").split("/")[-1].removesuffix(".git") or "project"


def _repo_name(url: str) -> str:
    name = _repo_slug(url).split("/")[-1]
    return re.sub(r"[^A-Za-z0-9_.-]+", "-", name).strip("-") or "project"


def _is_github_url(url: str) -> bool:
    return url.startswith("https://github.com/") or url.startswith("http://github.com/")


def _strip_github_proxy_prefix(url: str) -> str:
    for prefix in (
        "https://gh-proxy.org/",
        "http://gh-proxy.org/",
    ):
        if url.startswith(prefix):
            return url[len(prefix) :]
    return url


def _parse_score(text: str) -> int:
    patterns = [
        r"最终复现评分[：:]\s*(\d{1,3})",
        r"评分[：:]\s*(\d{1,3})",
        r"score\s*[=:]\s*(\d{1,3})",
        r"(\d{1,3})\s*/\s*100",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return max(0, min(100, int(match.group(1))))
    return 50 if text.strip() else 0


def _is_relative_to(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False


def to_jsonable(value: Any) -> Any:
    try:
        json.dumps(value)
        return value
    except TypeError:
        return str(value)
